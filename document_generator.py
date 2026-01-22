from docx import Document
from docx.shared import Inches

def create_word(text_blocks, image_paths, output_file):
    doc = Document()
    doc.add_heading("Extracted Chemistry Notes", level=1)

    for text in text_blocks:
        doc.add_paragraph(text)

    doc.add_page_break()
    doc.add_heading("Original Diagrams & Pages", level=2)

    for img in image_paths:
        doc.add_picture(img, width=Inches(5))

    doc.save(output_file)
